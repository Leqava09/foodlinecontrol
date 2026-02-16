# PDF Generation Examples for Django
# These alternatives work without LibreOffice on shared hosting

# ============================================
# 1. HTML to PDF (WeasyPrint - Best for styled documents)
# ============================================
from weasyprint import HTML
from django.template.loader import render_to_string
from django.http import HttpResponse

def generate_pdf_from_html(request):
    """Generate PDF from Django template"""
    # Render your template
    html_string = render_to_string('your_template.html', {'data': 'your_data'})
    
    # Convert to PDF
    html = HTML(string=html_string)
    pdf = html.write_pdf()
    
    # Return as response
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="report.pdf"'
    return response


# ============================================
# 2. DOCX to PDF (python-docx + reportlab)
# ============================================
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO

def docx_to_pdf_simple(docx_path, pdf_path):
    """Convert DOCX to PDF (basic text extraction)"""
    doc = Document(docx_path)
    
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    
    y = height - 50  # Start from top
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.strip():
            c.drawString(50, y, text)
            y -= 20
            if y < 50:  # New page if needed
                c.showPage()
                y = height - 50
    
    c.save()
    
    with open(pdf_path, 'wb') as f:
        f.write(buffer.getvalue())


# ============================================
# 3. XLSX to PDF (openpyxl + reportlab)
# ============================================
from openpyxl import load_workbook
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors

def xlsx_to_pdf(xlsx_path, pdf_path):
    """Convert Excel to PDF table"""
    wb = load_workbook(xlsx_path)
    ws = wb.active
    
    # Extract data
    data = []
    for row in ws.iter_rows(values_only=True):
        data.append(list(row))
    
    # Create PDF
    doc = SimpleDocTemplate(pdf_path, pagesize=landscape(letter))
    
    # Create table
    table = Table(data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    doc.build([table])


# ============================================
# 4. Create DOCX Reports (python-docx)
# ============================================
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report_docx(filename, data):
    """Create a DOCX report from scratch"""
    doc = Document()
    
    # Add heading
    heading = doc.add_heading('Production Report', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add paragraphs
    doc.add_paragraph(f"Report Date: {data['date']}")
    doc.add_paragraph(f"Batch Number: {data['batch']}")
    
    # Add table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Product'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Cost'
    
    # Data rows
    for product in data['products']:
        row_cells = table.add_row().cells
        row_cells[0].text = product['name']
        row_cells[1].text = str(product['qty'])
        row_cells[2].text = f"R {product['cost']}"
    
    # Save
    doc.save(filename)
    return filename


# ============================================
# 5. Create Reports from Scratch (ReportLab)
# ============================================
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib import colors

def create_advanced_pdf_report(filename, data):
    """Create professional PDF report with ReportLab"""
    doc = SimpleDocTemplate(filename, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f4788'),
        spaceAfter=30,
        alignment=1  # Center
    )
    story.append(Paragraph("Production Costing Report", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Introduction
    intro_text = f"<b>Date:</b> {data['date']}<br/><b>Site:</b> {data['site']}"
    story.append(Paragraph(intro_text, styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Table data
    table_data = [['Product', 'Units', 'Cost/Unit', 'Total']]
    for item in data['items']:
        table_data.append([
            item['name'],
            str(item['units']),
            f"R {item['cost_per_unit']:.2f}",
            f"R {item['total']:.2f}"
        ])
    
    # Create table
    table = Table(table_data, colWidths=[3*inch, 1.5*inch, 1.5*inch, 1.5*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f0f0f0')),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')])
    ]))
    
    story.append(table)
    
    # Build PDF
    doc.build(story)
    return filename


# ============================================
# 6. Django View Example - Invoice PDF
# ============================================
from django.views import View
from django.http import HttpResponse
from weasyprint import HTML
from django.template.loader import render_to_string

class InvoicePDFView(View):
    """Generate invoice PDF"""
    
    def get(self, request, invoice_id):
        # Get your data
        invoice = Invoice.objects.get(id=invoice_id)
        
        # Render HTML template
        html_string = render_to_string('invoices/invoice_pdf.html', {
            'invoice': invoice,
            'items': invoice.items.all(),
            'company': request.user.site.company_details
        })
        
        # Convert to PDF
        html = HTML(string=html_string)
        pdf = html.write_pdf()
        
        # Return as downloadable PDF
        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="invoice_{invoice_id}.pdf"'
        return response


# ============================================
# Usage in Django Admin
# ============================================
from django.contrib import admin
from django.urls import path
from django.http import HttpResponse

class YourModelAdmin(admin.ModelAdmin):
    
    def get_urls(self):
        urls = super().get_urls()
        custom_urls = [
            path('<int:object_id>/export-pdf/', self.export_pdf, name='export-pdf'),
        ]
        return custom_urls + urls
    
    def export_pdf(self, request, object_id):
        obj = self.get_object(request, object_id)
        
        # Generate PDF using WeasyPrint
        html_string = render_to_string('admin/export_template.html', {'object': obj})
        html = HTML(string=html_string)
        pdf = html.write_pdf()
        
        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{obj}_report.pdf"'
        return response
